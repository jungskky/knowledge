#!/usr/bin/env python3
"""
DOCX → PDF 변환, 페이지별 JPG 저장, 내부 이미지 추출 도구
============================================================
사용법:
    python docx_converter.py <input.docx> [옵션]

옵션:
    --pdf-only        PDF 변환만 수행 (JPG·이미지 추출 건너뜀)
    --images-only     DOCX 내부 이미지 추출만 수행
    --no-jpg          페이지 JPG 저장 건너뜀
    --dpi DPI         JPG 해상도 (기본값: 150)
    --output-dir DIR  출력 디렉토리 지정 (기본값: output/)
    --help            도움말 출력
"""

import sys
import os
import shutil
import zipfile
import argparse
import subprocess
import importlib.util
from pathlib import Path


# ─── soffice.py 로딩 ─────────────────────────────────────────────────────────
THIS_DIR = Path(__file__).resolve().parent

def get_soffice_runner():
    soffice_local = THIS_DIR / "soffice.py"
    if soffice_local.exists():
        spec = importlib.util.spec_from_file_location("soffice", soffice_local)
        mod  = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(mod)
        return mod.run_soffice

    # 폴백: 시스템 soffice 직접 호출
    def run_soffice(args):
        env = os.environ.copy()
        env["SAL_USE_VCLPLUGIN"] = "svp"
        return subprocess.run(["soffice"] + args, env=env, capture_output=True, text=True)
    return run_soffice


# ─── 1단계: DOCX → PDF ───────────────────────────────────────────────────────
def convert_to_pdf(docx_path: Path, output_dir: Path) -> Path:
    """DOCX를 PDF로 변환합니다."""
    print(f"\n[1/3] PDF 변환 중: {docx_path.name}")

    run_soffice = get_soffice_runner()
    result = run_soffice([
        "--headless",
        "--convert-to", "pdf",
        "--outdir", str(output_dir),
        str(docx_path)
    ])

    pdf_path = output_dir / (docx_path.stem + ".pdf")

    if result.returncode != 0:
        raise RuntimeError(f"PDF 변환 실패:\n{result.stderr}")
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF 파일이 생성되지 않았습니다: {pdf_path}")

    print(f"  ✅ {pdf_path.name}  ({pdf_path.stat().st_size / 1024:.1f} KB)")
    return pdf_path


# ─── 2단계: PDF → 페이지별 JPG ───────────────────────────────────────────────
def _find_pdftoppm() -> str | None:
    """pdftoppm 실행 파일 경로를 반환합니다 (없으면 None)."""
    import sys, shutil as sh
    if sys.platform == "win32":
        candidates = [
            r"C:\Program Files\poppler\Library\bin\pdftoppm.exe",
            r"C:\Program Files\poppler-xx\bin\pdftoppm.exe",
        ]
        for c in candidates:
            if Path(c).exists():
                return c
        return sh.which("pdftoppm")
    return sh.which("pdftoppm")


def pdf_to_jpg_pages(pdf_path: Path, output_dir: Path, dpi: int = 150) -> list[Path]:
    """
    PDF의 각 페이지를 JPG로 저장합니다.

    우선순위:
      1. pdf2image (Python 라이브러리, Pillow 기반)
      2. pdftoppm  (poppler 커맨드라인 도구)
    """
    print(f"\n[2/3] 페이지 JPG 변환 중 (DPI={dpi})")
    output_dir.mkdir(parents=True, exist_ok=True)
    stem = pdf_path.stem
    saved: list[Path] = []

    # ── 방법 1: pdf2image ──────────────────────────────────────────────────
    try:
        from pdf2image import convert_from_path
        pages = convert_from_path(str(pdf_path), dpi=dpi)
        for i, page in enumerate(pages, start=1):
            out = output_dir / f"{stem}_page{i:03d}.jpg"
            page.save(str(out), "JPEG", quality=92)
            saved.append(out)
            print(f"  🖼  {out.name}  ({out.stat().st_size / 1024:.1f} KB)")
        print(f"  ✅ 총 {len(saved)}페이지 저장 완료 → {output_dir}")
        return saved
    except ImportError:
        pass  # pdf2image 없으면 pdftoppm 시도
    except Exception as e:
        print(f"  ⚠️  pdf2image 오류: {e}  → pdftoppm 으로 재시도")

    # ── 방법 2: pdftoppm ──────────────────────────────────────────────────
    pdftoppm = _find_pdftoppm()
    if not pdftoppm:
        raise RuntimeError(
            "페이지 JPG 변환에 필요한 라이브러리가 없습니다.\n"
            "  pip install pdf2image Pillow\n"
            "또는 poppler 설치 후 다시 시도하세요.\n"
            "  Windows: https://github.com/oschwartz10612/poppler-windows/releases\n"
            "  Mac:     brew install poppler\n"
            "  Linux:   apt install poppler-utils"
        )

    prefix = str(output_dir / stem)
    subprocess.run(
        [pdftoppm, "-jpeg", "-r", str(dpi), str(pdf_path), prefix],
        check=True, capture_output=True
    )
    # pdftoppm 출력 파일을 _pageXXX.jpg 형식으로 정리
    raw_files = sorted(output_dir.glob(f"{stem}-*.jpg"))
    for i, f in enumerate(raw_files, start=1):
        renamed = output_dir / f"{stem}_page{i:03d}.jpg"
        f.rename(renamed)
        saved.append(renamed)
        print(f"  🖼  {renamed.name}  ({renamed.stat().st_size / 1024:.1f} KB)")

    print(f"  ✅ 총 {len(saved)}페이지 저장 완료 → {output_dir}")
    return saved


# ─── 3단계: DOCX 내부 이미지 추출 ───────────────────────────────────────────
SUPPORTED_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".emf", ".wmf", ".svg"}

def extract_embedded_images(docx_path: Path, output_dir: Path) -> list[Path]:
    """DOCX(ZIP) 내부 word/media/ 의 이미지를 추출합니다."""
    print(f"\n[3/3] 내부 이미지 추출 중: {docx_path.name}")

    img_dir = output_dir / f"{docx_path.stem}_images"
    img_dir.mkdir(parents=True, exist_ok=True)
    extracted: list[Path] = []

    with zipfile.ZipFile(docx_path, "r") as z:
        media_files = [
            n for n in z.namelist()
            if n.startswith("word/media/") and Path(n).suffix.lower() in SUPPORTED_EXTENSIONS
        ]
        if not media_files:
            print("  ℹ️  문서 내에 추출 가능한 이미지가 없습니다.")
            return extracted

        for member in media_files:
            dest = img_dir / Path(member).name
            # 파일명 충돌 방지
            counter = 1
            while dest.exists():
                dest = img_dir / f"{dest.stem}_{counter}{dest.suffix}"
                counter += 1

            with z.open(member) as src, open(dest, "wb") as dst:
                shutil.copyfileobj(src, dst)
            extracted.append(dest)
            print(f"  📷 {dest.name}  ({dest.stat().st_size / 1024:.1f} KB)")

    print(f"  ✅ 총 {len(extracted)}개 이미지 추출 → {img_dir}")
    return extracted


# ─── 문서 정보 출력 ───────────────────────────────────────────────────────────
def print_doc_info(docx_path: Path):
    try:
        from docx import Document
        doc = Document(docx_path)
        print(f"\n📄 문서 정보: {docx_path.name}")
        print(f"   단락 수   : {len(doc.paragraphs)}")
        print(f"   표 수     : {len(doc.tables)}")
        if doc.core_properties.author:
            print(f"   작성자    : {doc.core_properties.author}")
        if doc.core_properties.title:
            print(f"   제목      : {doc.core_properties.title}")
    except Exception:
        pass


# ─── CLI ─────────────────────────────────────────────────────────────────────
def parse_args():
    parser = argparse.ArgumentParser(
        description="DOCX → PDF 변환, 페이지 JPG 저장, 내부 이미지 추출",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("input",          help="변환할 .docx 파일 경로")
    parser.add_argument("--pdf-only",     action="store_true", help="PDF 변환만 수행")
    parser.add_argument("--images-only",  action="store_true", help="DOCX 내부 이미지 추출만 수행")
    parser.add_argument("--no-jpg",       action="store_true", help="페이지 JPG 저장 건너뜀")
    parser.add_argument("--dpi",          type=int, default=150, help="JPG 해상도 (기본: 150)")
    parser.add_argument("--output-dir",   "-o", default="output", help="출력 디렉토리 (기본: output/)")
    return parser.parse_args()


def main():
    args = parse_args()

    docx_path = Path(args.input).resolve()
    if not docx_path.exists():
        print(f"❌ 파일을 찾을 수 없습니다: {docx_path}")
        sys.exit(1)
    if docx_path.suffix.lower() != ".docx":
        print(f"❌ .docx 파일만 지원합니다.")
        sys.exit(1)

    output_dir = Path(args.output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    print("=" * 55)
    print("  DOCX 변환 도구")
    print("=" * 55)
    print_doc_info(docx_path)

    results = {}

    try:
        # 1. PDF 변환 (images-only가 아닌 경우)
        if not args.images_only:
            pdf_path = convert_to_pdf(docx_path, output_dir)
            results["pdf"] = pdf_path

            # 2. 페이지별 JPG (pdf-only, no-jpg가 아닌 경우)
            if not args.pdf_only and not args.no_jpg:
                pages = pdf_to_jpg_pages(pdf_path, output_dir, dpi=args.dpi)
                results["pages"] = pages

        # 3. DOCX 내부 이미지 추출 (pdf-only가 아닌 경우)
        if not args.pdf_only:
            embedded = extract_embedded_images(docx_path, output_dir)
            results["embedded"] = embedded

    except Exception as e:
        print(f"\n❌ 오류 발생: {e}")
        sys.exit(1)

    # ─── 최종 요약 ────────────────────────────────────────────────────────────
    print("\n" + "=" * 55)
    print("  처리 완료 요약")
    print("=" * 55)
    if "pdf" in results:
        print(f"  📄 PDF         : {results['pdf'].name}")
    if "pages" in results:
        print(f"  🖼  페이지 JPG  : {len(results['pages'])}장 → {output_dir}")
    if "embedded" in results:
        count = len(results["embedded"])
        if count:
            print(f"  📷 내부 이미지 : {count}개 → {output_dir / (docx_path.stem + '_images')}")
        else:
            print(f"  📷 내부 이미지 : 없음")
    print(f"  📁 출력 폴더   : {output_dir}")
    print()


if __name__ == "__main__":
    main()
